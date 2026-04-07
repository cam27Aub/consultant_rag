from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os, uuid, tempfile

app = FastAPI()

class ReportRequest(BaseModel):
    markdown: str
    filename: str = "ConsultantIQ_Report"

def clean_signal_tags(text: str) -> str:
    return text.replace("%%REPORT_READY%%", "").replace("%%DECK_READY%%", "").strip()

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3564')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def set_heading_style(para, level):
    for run in para.runs:
        if level == 1:
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
        elif level == 2:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.name = 'Calibri'

def add_table_from_rows(doc, table_rows):
    # Filter out separator rows like |---|---|
    data_rows = [r for r in table_rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
    if not data_rows:
        return

    parsed = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cells)

    if not parsed:
        return

    max_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=1, cols=max_cols)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, val in enumerate(parsed[0]):
        if i < max_cols:
            hdr_cells[i].text = val
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
            # Dark header background
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1F3564')
            tcPr.append(shd)

    # Data rows
    for row_data in parsed[1:]:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            if i < max_cols:
                row[i].text = val
                for run in row[i].paragraphs[0].runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)

    doc.add_paragraph()

def markdown_to_docx(markdown: str, filename: str) -> str:
    markdown = clean_signal_tags(markdown)
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default paragraph style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    lines = markdown.split('\n')
    i = 0
    table_buffer = []

    while i < len(lines):
        line = lines[i]

        # Collect table rows
        if line.strip().startswith('|'):
            table_buffer.append(line)
            i += 1
            continue
        else:
            if table_buffer:
                add_table_from_rows(doc, table_buffer)
                table_buffer = []

        # H1
        if line.startswith('# '):
            para = doc.add_heading(line[2:].strip(), level=1)
            set_heading_style(para, 1)

        # H2
        elif line.startswith('## '):
            para = doc.add_heading(line[3:].strip(), level=2)
            set_heading_style(para, 2)

        # H3
        elif line.startswith('### '):
            para = doc.add_heading(line[4:].strip(), level=3)
            set_heading_style(para, 3)

        # Horizontal rule
        elif line.strip() == '---':
            add_horizontal_rule(doc)

        # Bullet point
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:].strip()
            # Handle inline bold in bullets
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # Bold key-value metadata lines e.g. **Prepared by:** value
        elif re.match(r'^\*\*(.+?):\*\*\s*(.*)', line):
            match = re.match(r'^\*\*(.+?):\*\*\s*(.*)', line)
            p = doc.add_paragraph()
            run_key = p.add_run(match.group(1) + ': ')
            run_key.bold = True
            run_key.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
            run_key.font.name = 'Calibri'
            p.add_run(match.group(2))

        # Empty line — spacing
        elif line.strip() == '':
            doc.add_paragraph()

        # Normal paragraph with possible inline bold
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.name = 'Calibri'
                else:
                    run = p.add_run(part)
                    run.font.name = 'Calibri'

        i += 1

    # Flush any remaining table
    if table_buffer:
        add_table_from_rows(doc, table_buffer)

    # Save to temp
    tmp_dir = tempfile.gettempdir()
    safe_name = re.sub(r'[^\w\-]', '_', filename)
    filepath = os.path.join(tmp_dir, f"{safe_name}_{uuid.uuid4().hex[:8]}.docx")
    doc.save(filepath)
    return filepath


@app.post("/generate-report")
async def generate_report(req: ReportRequest):
    try:
        filepath = markdown_to_docx(req.markdown, req.filename)
        return FileResponse(
            path=filepath,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=os.path.basename(filepath),
            headers={"Content-Disposition": f"attachment; filename={os.path.basename(filepath)}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
